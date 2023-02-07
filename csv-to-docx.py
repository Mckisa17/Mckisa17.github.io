

import pandas as pd
import datetime as dt
import docx
from docx import Document
from docx.shared import Pt
from docx.shared import Inches

from flask import Flask, render_template, request
import pandas as pd

app = Flask(__name__)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])

# Read the CSV file into a pandas dataframe
df = pd.read_csv('/kaggle/input/jit-csv/JIT.csv')

# Extract the values
df = pd.DataFrame(df)

df.loc[df["Secondary?"] == "N", "Secondary?"] = "Institution"
df['date_column'] = pd.to_datetime(df['Project Period End Date'], format='%m/%d/%Y')



def create_word_document():
    # Create a new Word document
    doc = docx.Document()

    # Set font and size
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    doc.styles['Normal'].paragraph_format.line_spacing = 1.0
    doc.styles['Normal'].paragraph_format.space_after = Pt(0)

    # Loop over grant_ID values
    for i in range(df.shape[0]):
        
        doc.add_paragraph("*Title: " + df["Grant Title"].loc[i])

        # Insert the rest of the text
        doc.add_paragraph("Major Goals: ")
        doc.add_paragraph("*Status of Support: "+ df["Status Category"].loc[i])
        doc.add_paragraph("Project Number: " + df["External Grant ID"].astype(str).loc[i])
        doc.add_paragraph("Name of PD/PI: " + df["PI (last, first)"].loc[i])
        doc.add_paragraph("*Source of Support: "+ df["Funding Agency"].loc[i])
        doc.add_paragraph("*Primary Place of Performance: "+ df["Secondary?"].loc[i])
        doc.add_paragraph("Project/Proposal Start and End Date: (MM/YYYY) (if available): "+ df["Project Period Start Date"].loc[i] + " - " +df["Project Period End Date"].loc[i])
        doc.add_paragraph("* Total Award Amount (including Indirect Costs): "+ df["Total Project Period Costs (Direct plus Indirect)"].loc[i])
        doc.add_paragraph("* Person Months (Calendar/Academic/Summer) per budget period.")
        doc.add_paragraph()
            
        # Add a table to the document
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Table Grid'
        cell = table.cell(0, 0)
        cell.text = "Year (YYYY)"
        cell = table.cell(0, 1)
        cell.text = "Person Months (##.##)"
        doc.add_paragraph()

        # Set the cell height and width
        for row in table.rows:
            for cell in row.cells:
                cell.width = Inches(1.81)
                cell.height = Inches(0.18)
                cell_2 = table.cell(i, 1)
                cell_2.text = df['Period 1 Effort (Calendar months)'].astype(str).loc[i]
        
        df['year'] = df['date_column'].dt.year
        table.cell(1, 0).text = "1. "  + str(df['year'].astype(int).loc[i] - 4)
        table.cell(1, 1).text = df['Period 1 Effort (Calendar months)'].astype(str).loc[i]
        table.cell(2, 0).text = "2. "  + str(df['year'].astype(int).loc[i] - 3)
        table.cell(2, 1).text = df['Period 2 Effort (Calendar months)'].astype(str).loc[i]
        table.cell(3, 0).text = "3. "  + str(df['year'].astype(int).loc[i] - 2)
        table.cell(3, 1).text = df['Period 3 Effort (Calendar months)'].astype(str).loc[i]
        table.cell(4, 0).text = "4. "  + str(df['year'].astype(int).loc[i] - 1)
        table.cell(4, 1).text = df['Period 4 Effort (Calendar months)'].astype(str).loc[i]
        table.cell(5, 0).text = "5. "  + str(df['year'].astype(int).loc[i])
        table.cell(5, 1).text = df['Period 5 Effort (Calendar months)'].astype(str).loc[i]
        

                

    # Save the document
    doc.save('output1.docx')

create_word_document()
